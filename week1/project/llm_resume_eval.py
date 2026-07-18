import os
import time
import csv  # <-- Added for CSV export
import json
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client = Groq(api_key=my_api_key)
model = "llama-3.3-70b-versatile"

job_description = """
Job Requisition ID #

26WD99919
Position Overview

Autodesk's Identity and Access Management team is seeking a talented and highly motivated Software Engineer who enjoys learning, solving complex problems, and building reliable software solutions.

As a member of the team responsible for foundational identity services, you will contribute to Autodesk's Identity and Access Management platform, which supports millions of Autodesk customers and powers more than 150 Autodesk desktop, mobile, and web applications, cloud services, and third-party developer integrations.

This is an excellent opportunity for an early-career engineer to work with modern cloud technologies, collaborate with experienced engineers, and gain hands-on experience building scalable, secure, and reliable services that are critical to Autodesk's products and customers.

Responsibilities

Work as a member of a self-organized Agile team that builds, owns, and operates cloud services
Contribute to the development of cloud services, including back-end, front-end, DevOps, and quality engineering activities
Support service operations, including monitoring, alerting, logging, metrics, and troubleshooting
Collaborate closely with senior engineers, architects, and product owners to understand requirements and implement effective solutions
Write clean, maintainable, and well-tested code under the guidance of experienced team members
Participate in code reviews, design discussions, sprint planning, and other Agile ceremonies
Learn and apply engineering best practices for scalability, reliability, usability, performance, and security
Support the continuous improvement of existing systems by identifying issues and recommending enhancements
Minimum Qualifications

Bachelor's degree or equivalent experience in Computer Science, Software Engineering, or a related technical field
0–1 year of software engineering experience, including internships, academic projects, or professional experience
Good understanding of fundamental computer science concepts and programming principles
Hands-on experience with at least one programming language, such as C#, Go, Java, or Python
Familiarity with web services, Representational State Transfer (REST) Application Programming Interfaces (APIs), and client-server architecture
Basic understanding of databases, data structures, algorithms, and object-oriented programming
Willingness to learn cloud technologies, DevOps practices, and production service operations
Strong problem-solving, communication, and collaboration skills
Preferred Qualifications

Internship, academic, or project experience developing web applications, back-end services, or Application Programming Interfaces (APIs)
Basic exposure to cloud platforms such as Amazon Web Services (AWS), Microsoft Azure, or Google Cloud Platform (GCP)
Familiarity with Git, Continuous Integration and Continuous Deployment (CI/CD) pipelines, unit testing, and debugging tools
Exposure to authentication, authorization, identity and access management, or application security concepts
Ability to learn quickly and contribute effectively in a collaborative engineering environment
The Ideal Candidate

The ideal candidate is an enthusiastic early-career software engineer who is passionate about learning, solving technical challenges, and building reliable software that delivers value to customers

Demonstrates a strong foundation in computer science principles and software engineering fundamentals
Enjoys solving technical problems through logical thinking, curiosity, and analytical reasoning
Takes ownership of assigned work while actively seeking opportunities to learn and grow
Writes clean, maintainable, and high-quality code while embracing engineering best practices
Collaborates effectively with engineers, product managers, architects, and other cross-functional partners
Welcomes feedback and applies it to continuously improve technical and professional skills
Demonstrates a willingness to learn new technologies, cloud platforms, and modern software development practices
Approaches challenges with a growth mindset and adapts quickly in a fast-paced engineering environment
Communicates technical ideas clearly and works effectively within a collaborative team
Builds strong relationships with teammates and contributes positively to the team's culture and success
Demonstrates attention to detail while understanding how individual contributions support broader product and business goals

"""

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system = {
    "role": "system",
    "content": system_prompt
}
message_user = {
    "role": "user",
    "content": user_prompt
}
response_format = {
    "type": "json_object"
}

messages = [message_system, message_user]

response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
answer = response.choices[0].message.content
raw_json = answer

job_data = json.loads(raw_json)
job = JobD(**job_data)

print("Min Experience Required:", job.minimum_experience)
print("Edu Requirements:", job.education_requirements)

# parse real
class MatchResult(BaseModel):
    score: float
    details: dict

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()

def final_score(job, resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message = {
        "role": "user",
        "content": prompt
    }
    messages = [message]
    response_format = {
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system = {
        "role": "system",
        "content": system_prompt
    }
    message_user = {
        "role": "user",
        "content": user_prompt
    }
    messages = [message_system, message_user]
    response_format = {
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None

# lets do it now
resume_folder = Path("resumes")
all_results = []
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)  # llm call 1
    time.sleep(5)
    result = final_score(job, parsed_resume)  # llm call 2
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name or file_path.stem,
        "score": result.score,
        "details": result.details
    })

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)

top_2 = all_results[:2]
worst_2 = all_results[-2:]

print("\nTOP 2 CANDIDATES")
for candidate in top_2:
    print(candidate["name"], "-", candidate["score"], "%")
    print(candidate["details"])

print("\nLOWEST 2 CANDIDATES")
for candidate in worst_2:
    print(candidate["name"], "-", candidate["score"], "%")
    print(candidate["details"])

# ==========================================
# NEW: EXPORT RESULT TO CSV
# ==========================================
csv_filename = "evaluation_report.csv"
print(f"\nExporting results to {csv_filename}...")

with open(csv_filename, mode="w", newline="", encoding="utf-8") as file:
    writer = csv.writer(file)
    # Write CSV Header
    writer.writerow([
        "Rank", 
        "Candidate Name", 
        "Match Score (%)", 
        "Matching Skills", 
        "Missing Important Skills", 
        "Experience Requirements Met", 
        "Final Verdict"
    ])
    
    # Write data rows
    for index, candidate in enumerate(all_results, 1):
        details = candidate["details"]
        
        # Safely extract details keys in case LLM labels keys slightly differently
        matching_skills = details.get("matching_skills", details.get("2", ""))
        missing_skills = details.get("missing_important_skills", details.get("missing_skills", details.get("3", "")))
        experience_met = details.get("experience_requirement_met", details.get("4", ""))
        verdict = details.get("short_final_verdict", details.get("verdict", details.get("6", "")))
        
        # Convert lists to comma-separated strings if returned as list objects
        if isinstance(matching_skills, list):
            matching_skills = ", ".join(matching_skills)
        if isinstance(missing_skills, list):
            missing_skills = ", ".join(missing_skills)
            
        writer.writerow([
            index,
            candidate["name"],
            f"{candidate['score']}%",
            matching_skills,
            missing_skills,
            experience_met,
            verdict
        ])

print("Export complete!")